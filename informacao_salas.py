
import os
import csv
from docx import Document
import cv2


if __name__ == "__main__":

    documento = Document("modelo_informacao.docx")



    with open(os.path.join(os.getcwd(), "Organização das salas.csv"), "r") as arquivo:
        arquivo_csv = csv.reader(arquivo, delimiter = ";")

        for linha in arquivo_csv:

            sala = str(linha[1])
            projetor = str(linha[2])
            patrimonio = str(linha[3])
            capacidade = str(linha[4])

            for paragrafo in documento.paragraphs:

                referencia = {
                "s" : sala ,
                "pj" : projetor,
                "pt" : patrimonio,
                "cp" : capacidade,
            }
                for paragrafo in documento.paragraphs:
                    for codigo in referencia:
                        valor = referencia[codigo]
                        paragrafo.text = paragrafo.text.replace(codigo, valor)

                documento.save(f"Sala - {sala}.docx")


        

                    # if "x" in paragrafo.text:
                    #     paragrafo.text = paragrafo.text.replace("x", sala)
                    # elif "xx" in paragrafo.text:
                    #     paragrafo.text = paragrafo.text.replace("xx", projetor)
                    # elif "xxx" in paragrafo.text:
                    #     paragrafo.text = paragrafo.text.replace("xxx", patrimonio)
                    # elif "xxxx" in paragrafo.text:
                    #     paragrafo.text = paragrafo.text.replace("xxxx", capacidade)